from carta import *
import docx
import serial
#primeira rodada

r1 = Rodada_1()


arquetipoID = ['92:BB:2E:00', '2', '3']
animalID = ['93:DB:B9:79', '5', '6']
metaforaID = ['D5:32:92:03', '8', '9', '10', '11', '12', '13']

lista =[]

ser = serial.Serial('/dev/cu.usbmodem14301', 9600)

def abre_serial():
	if ser.is_open:
		ser.close()
		ser.open()
	else:
		ser.open()

def read_ser():
	linha = ser.readLine()
	return linha
#quando o serial.print tiver um /n ele vai ler e retornar

doc = docx.Document()

#Cada carta poderia ter um único ID. Assim, será fácil de identificar as cartas pelo arduino, já que estarão

#Precisamos solucionar esse problema do documento.

nome = input('Digite um nome para o documento: ') + '.docx'

print('--------------------------------------------------------------------------')
print('|                    |    Bem-vindo(a) ao ABSURD!    |                   |')
print('--------------------------------------------------------------------------')

if ser.readline() == "92:BB:2E:00":
	arquetipo = "O Monstro"

if ser.readline() == 'D5:32:92:03':
	status_quo = 'Perdido na Multidão'

if ser.readline() == '93:DB:B9:79':
	animal = 'Black Elephant'

while arquetipoID.count(arquetipo) == 0:
    print("Carta Errada! Escolha a Carta certa para continuar.")
    arquetipo
lista.append(arquetipo)

while metaforaID.count(status_quo) == 0:
    print("Carta Errada! Escolha a Carta certa para continuar.")
    status_quo
lista.append(status_quo)

while animalID.count(animal) == 0:
    print("Carta Errada! Escolha a Carta certa para continuar.")
    animal
lista.append(animal)

r1.append_a(lista[0])
r1.append_sq(lista[1])
r1.append_an(lista[2])

doc.add_paragraph('O jogo começou! primeiramente os jogadores escolheram o arquétipo: ' + arquetipo)

doc.add_paragraph('O animal da rodada foi ' + animal)

doc.add_paragraph('Logo após a escolha do arquétipo, os jogadores continuaram e escolheram as 5 seguintes cartas: ')
#...
doc.add_paragraph("Primeira Rodada: ", r1.print_states())



print("Ótimo! seu documento foi criado! é só procurar pelo arquivo " + nome)

print('--------------------------------------------------------------------------')

doc.save(nome)